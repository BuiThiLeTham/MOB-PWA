from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = r"C:\Users\31bui\Downloads\VKU-Field-Survey-PWA-Technical-Report.docx"


def set_page_number_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_paragraph(doc, text, bold=False, italic=False, style=None, align=None, size=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
      p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size is not None:
        r.font.size = Pt(size)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
set_page_number_footer(section)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(12)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("MINI-PROJECT SHORT TECHNICAL REPORT")
run.bold = True
run.font.size = Pt(14)

add_paragraph(doc, "Course: Cross-Platform Mobile App Development (VKU)")
add_paragraph(doc, "Mini-Project Title: VKU Field Survey PWA")
add_paragraph(doc, "Team / Student Name: [Please fill in]")
add_paragraph(doc, "Submission Date: 03/09/2026")

add_paragraph(doc, "")
add_paragraph(doc, "1. GENERAL INFORMATION & DELIVERABLE LINKS", bold=True)
add_paragraph(doc, "Team Members:", bold=True)
add_paragraph(doc, "1. [Please fill in] - Student ID: [Please fill in] - Role: Full-stack/PWA - Contribution: 100%")
add_paragraph(doc, "Live Demo URL: https://vku-field-survey-pwa.pages.dev")
add_paragraph(doc, "GitHub Repository: https://github.com/BuiThiLeTham/MOB-PWA")
add_paragraph(doc, "Video Demo: [Please fill in if available].")

add_paragraph(doc, "")
add_paragraph(doc, "2. FEATURE IMPLEMENTATION CHECKLIST", bold=True)
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "No."
hdr[1].text = "Required Feature"
hdr[2].text = "Status"
hdr[3].text = "Implementation Details & Acceptance Level"

rows = [
    ("1", "Responsive mobile viewport", "Complete", "Responsive HTML/CSS layout optimized for mobile, tablet, and desktop with app-like navigation."),
    ("2", "Web App Manifest (installable)", "Complete", "manifest.json defines standalone display, theme colors, start_url, scope, and install icons including maskable variants."),
    ("3", "Service Worker lifecycle", "Complete", "service-worker.js implements install, activate, fetch, and background sync event handling."),
    ("4", "Caching strategies", "Complete", "Cache-First for app shell, Network-First for surveys, and Network-Only for API POST requests."),
    ("5", "Local offline persistence", "Complete", "IndexedDB FieldSurveyDB stores surveys, responses, syncQueue, and settings on the client device."),
    ("6", "Offline draft / submit queue", "Complete", "Responses are saved locally first with pending/syncing/synced/failed states before any network call."),
    ("7", "Automatic background sync", "Complete", "Background Sync uses tag sync-responses; browsers without support fall back to the online event retry flow."),
    ("8", "HTTPS live deploy", "Complete", "Cloudflare Pages production deployment is available at https://vku-field-survey-pwa.pages.dev."),
]
for no, feature, status, detail in rows:
    cells = table.add_row().cells
    cells[0].text = no
    cells[1].text = feature
    cells[2].text = status
    cells[3].text = detail

add_paragraph(doc, "")
add_paragraph(doc, "3. TECHNICAL ARCHITECTURE & PROJECT STRUCTURE", bold=True)
add_paragraph(
    doc,
    "The client is a plain HTML/CSS/JavaScript Progressive Web App. The main application flow is managed by js/app.js, which coordinates routing, rendering, submission, settings, install prompts, and synchronization status.",
)
add_paragraph(doc, "Directory map:", bold=True)
for line in [
    "index.html - application shell and entry point",
    "manifest.json - PWA install metadata and icons",
    "service-worker.js - precache, fetch strategies, and background sync",
    "js/app.js - app controller, routes, submit flow, settings, install prompt",
    "js/db.js - IndexedDB schema and persistence helpers",
    "js/sync.js - sync manager and online-event fallback",
    "js/api.js - Google Apps Script API access",
    "js/survey.js - survey loading, local caching, and data preparation",
    "js/form.js - dynamic form rendering and validation",
    "data/surveys.json - bundled survey definitions for offline use",
    "apps-script/Code.gs - backend Web App logic for Sheets read/write and duplicate protection",
    "scripts/build-pages.mjs - builds the dist folder for Cloudflare Pages",
]:
    add_paragraph(doc, line, style="List Bullet")

add_paragraph(
    doc,
    "State flow: user selects a survey -> fills the dynamic form -> client validates inputs -> response is saved to IndexedDB -> if online the sync manager posts to Google Apps Script -> if offline the response remains pending and is retried later. Duplicate writes are prevented server-side using responseId.",
)
add_paragraph(
    doc,
    "Storage model: IndexedDB contains four object stores: surveys, responses, syncQueue, and settings. This allows the app to preserve survey definitions, response history, device configuration, and unsent records even after refresh or browser restart.",
)
add_paragraph(
    doc,
    "Deployment model: the frontend is hosted as a static PWA on Cloudflare Pages over HTTPS, while Google Apps Script acts as the lightweight backend adapter that writes records into Google Sheets.",
)

add_paragraph(doc, "")
add_paragraph(doc, "4. EMPIRICAL EVIDENCE & SCREENSHOTS", bold=True)
add_paragraph(doc, "Suggested screenshots from the live URL on a mobile viewport. Replace the placeholders below with your actual captures before submission.")
for caption in [
    "Figure 1. Home screen: survey list, online/offline badge, and pending sync counter.",
    "Figure 2. Survey form: dynamic question rendering with validation messages.",
    "Figure 3. History or Sync screen: local responses with pending/synced states.",
    "Figure 4. Installed PWA or Settings screen: Apps Script URL, app mode, and configuration.",
]:
    add_paragraph(doc, caption, bold=True)
    add_paragraph(doc, "[Insert screenshot here]", italic=True)

add_paragraph(
    doc,
    "Additional demo evidence can include IndexedDB records in DevTools and a matching row added to Google Sheets after connectivity is restored.",
)

add_paragraph(doc, "")
add_paragraph(doc, "5. TECHNICAL CHALLENGES & RESOLUTIONS", bold=True)
add_paragraph(doc, "5.1 Preventing data loss in unstable network conditions", bold=True)
add_paragraph(
    doc,
    "Field survey users may submit forms in weak or missing connectivity. To avoid data loss, the app always writes responses into IndexedDB first and only marks them as synced after the server returns a successful response. If the network is unavailable, the record stays on the device with pending or failed status and can be retried later.",
)
add_paragraph(doc, "5.2 Avoiding duplicate records during retry", bold=True)
add_paragraph(
    doc,
    "Automatic retry can accidentally create duplicate rows on the backend. This project solves the problem by attaching a unique responseId to every submission and checking that ID in Google Apps Script before appending to the Responses sheet. Re-sending the same record therefore updates status without creating duplicate entries.",
)
add_paragraph(doc, "5.3 Browser differences in background sync support", bold=True)
add_paragraph(
    doc,
    "Not all browsers implement Background Sync consistently. The service worker attempts to register sync-responses when supported, while the client also listens for the online event as a fallback. This keeps synchronization reliable across a wider range of browsers and devices.",
)
add_paragraph(doc, "5.4 CORS and Apps Script request compatibility", bold=True)
add_paragraph(
    doc,
    "Google Apps Script web apps can be sensitive to preflight requests and content-type handling. To reduce CORS friction, the client sends text/plain payloads instead of standard JSON posts, which keeps the integration simpler and more reliable for this mini-project deployment.",
)

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
