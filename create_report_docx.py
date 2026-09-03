from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = r"C:\Users\31bui\Downloads\Mini-Project-1-Technical-Report.docx"


def set_page_number_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_paragraph(doc, text, bold=False, italic=False, style=None, align=None, size=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size is not None:
        r.font.size = Pt(size)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
set_page_number_footer(section)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(12)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("MINI-PROJECT SHORT TECHNICAL REPORT")
run.bold = True
run.font.size = Pt(14)

add_paragraph(doc, "Course: Cross-Platform Mobile App Development (VKU)")
add_paragraph(doc, "Mini-Project Title: Mini-Project 1 - Week 3 PWA (Trovey Field Survey)")
add_paragraph(doc, "Team / Student Name: Nguyen Minh Duy")
add_paragraph(doc, "Submission Date: 03/09/2026")

add_paragraph(doc, "")
add_paragraph(doc, "1. GENERAL INFORMATION & DELIVERABLE LINKS", bold=True)
add_paragraph(doc, "Team Members:", bold=True)
add_paragraph(
    doc,
    "1. Nguyen Minh Duy - Student ID: 23IT038 - Role: Solo (UI, PWA, sync, deploy) - Contribution: 100%",
)
add_paragraph(doc, "Live Demo URL: https://app.puretrovey.net/")
add_paragraph(doc, "GitHub Repository: https://github.com/minhduy6868/MOB-Survey")
add_paragraph(doc, "Video Demo: Not submitted.")

add_paragraph(doc, "")
add_paragraph(doc, "2. FEATURE IMPLEMENTATION CHECKLIST", bold=True)

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "No."
hdr[1].text = "Required Feature"
hdr[2].text = "Status"
hdr[3].text = "Implementation Details & Acceptance Level"

rows = [
    ("1", "Responsive mobile viewport", "Complete", "Flutter web, Material 3 NavigationBar, 360-430 px. Light/dark in Settings."),
    ("2", "Web App Manifest (installable)", "Complete", "app/web/manifest.json: display standalone, start_url/scope HTTPS, icons 192/512 + maskable."),
    ("3", "Service Worker lifecycle", "Complete", "app/web/sw.js: install (precache), activate (drop old caches), fetch. Cache name trovey-shell-v9."),
    ("4", "Five caching strategies", "Complete", "Cache-First shell; Network-First /survey-template.json; SWR /sw-stats.json; Cache-Only /offline; Network-Only /api/*."),
    ("5", "Local offline persistence", "Complete", "Hive boxes tickets-v3, settings, meta. On web this is IndexedDB."),
    ("6", "Offline draft + submit queue", "Complete", "Draft saved in Hive. Submit sets queued. Status: draft | queued | syncing | synced | failed."),
    ("7", "Automatic background sync", "Complete", "connectivity_plus drain on reconnect. SW sync tag trovey-sync posts DRAIN_QUEUE to clients."),
    ("8", "HTTPS live deploy", "Complete", "Cloudflare Pages. Canonical host https://app.puretrovey.net/."),
]

for no, feature, status, detail in rows:
    cells = table.add_row().cells
    cells[0].text = no
    cells[1].text = feature
    cells[2].text = status
    cells[3].text = detail

add_paragraph(doc, "")
add_paragraph(doc, "3. TECHNICAL ARCHITECTURE & PROJECT STRUCTURE", bold=True)
add_paragraph(
    doc,
    "Client is Flutter web. TroveyStore (ChangeNotifier) is the only app state. It owns Hive boxes, connectivity, and the sync queue. Screens read the store and call put / enqueue / drainQueue.",
)
add_paragraph(doc, "Directory map:", bold=True)
for line in [
    "app/lib/data/store.dart - Hive + queue (tickets-v3)",
    "app/lib/data/sync.dart - POST /api/sync, GET /api/records",
    "app/lib/screens/ - Home, form (3 steps), results, settings",
    "app/web/manifest.json - Install metadata",
    "app/web/sw.js - SW install / activate / fetch + sync",
    "functions/api/sync.js - Write Cloudflare KV (TROVEY_RECORDS)",
    "functions/api/records.js - JSON / CSV read API",
    "sheets/Code.gs - Sheet pulls KV; not the primary store",
]:
    add_paragraph(doc, line, style="List Bullet")

add_paragraph(
    doc,
    "State flow: form edit -> Hive draft -> Submit -> status queued -> SyncClient.send -> KV put ticket:{clientId}. clientId is a UUID, so retry does not create a second row. On online, pullCloud() merges remote rows, then drainQueue() sends local queued/failed rows.",
)
add_paragraph(
    doc,
    "Exception handling: HTTP or body ok!=true / kv!=true throws. sendOne() sets status failed and lastError. pullCloud() swallows network errors and keeps the local box. The UI shows queued / failed / retry. A ticket is accepted when KV write succeeds. Sheet webhook errors do not fail the ticket.",
)
add_paragraph(
    doc,
    "Cache map used in the report demo: navigations try network then fall back to cached /index.html or /offline; API is never cached.",
)

add_paragraph(doc, "")
add_paragraph(doc, "4. EMPIRICAL EVIDENCE & SCREENSHOTS", bold=True)
add_paragraph(doc, "Screenshots from the live URL on a 390x844 viewport (Chromium). Captured 03/09/2026.")
for caption in [
    "Figure 1. Home: topic, start CTA, submitted/queue counts, install card.",
    "Figure 2. Form step 1 (site): country dropdown and required place field.",
    "Figure 3. Results: local totals after submit.",
    "Figure 4. Settings: collector 23IT038 (read-only), locale, theme.",
]:
    add_paragraph(doc, caption, bold=True)
    add_paragraph(doc, "[Insert screenshot here]", italic=True)

add_paragraph(
    doc,
    "Cache-Only fallback /offline is in the live host (Phieu van o tren may). It is used when a navigation miss is not in cache.",
)

add_paragraph(doc, "")
add_paragraph(doc, "5. TECHNICAL CHALLENGES & RESOLUTIONS", bold=True)
add_paragraph(doc, "5.1 Chrome ERR_FAILED on the apex host", bold=True)
add_paragraph(
    doc,
    "https://puretrovey.net/ returned 302 to the app host, but some campus Chrome sessions showed ERR_FAILED. The origin was up. The browser reused a cached Alt-Svc HTTP/3 (h3) route that did not complete on that network. Change: document https://app.puretrovey.net/ as the demo URL; send Alt-Svc: clear on / and /index.html; service worker uses network-first for navigations. Workaround for an old home-screen icon: open the app host or https://app.puretrovey.net/?reset=1.",
)
add_paragraph(doc, "5.2 Google Sheets webhook returned 401", bold=True)
add_paragraph(
    doc,
    "Apps Script doPost required a signed-in user (Who has access != Anyone), so the webhook could not be the database. Change: Cloudflare KV is the store. POST /api/sync writes KV and returns ok only if kv is true. The sheet runs beautifyAndSync() / UrlFetchApp on GET /api/records. Sheet downtime does not block submit.",
)

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
